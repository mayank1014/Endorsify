from Listofthings import PronounsList, PositivePersonalityTraits, AcademicSkills, Phrase1, Phrase2, Phrase3, Phrase4, Phrase5, LinkingWords
import sys
import os
import json
import random
import datetime
from datetime import date
import sys
import json
import cloudinary.uploader
from docx import Document
import requests 
from io import BytesIO 
from docx.shared import Inches


# Access command-line argument containing the JSON string
json_string = sys.argv[1]

# Parse the JSON string into a Python object
json_object = json.loads(json_string)

# Access the JSON object's properties
for key, value in json_object.items():
    globals()[key] = value
    


response = requests.get(docxFile)
docx_bytes = BytesIO(response.content)
doc = Document(docx_bytes)



# doc = docx.Document()



all_paras = doc.paragraphs



pronounLocalList = PronounsList.get(pronoun) 
# he
subjective = pronounLocalList[0] 
# him
objective = pronounLocalList[1] 
# his
possessive = pronounLocalList[2] 



############################
# Time teacher know students
############################

yearsAttended = schoolYearAttended.split('/') 
start_date = datetime.datetime(int(yearsAttended[0]), 8, 15) 
num_months = (date.today().year - start_date.year) * 12 + (date.today().month - start_date.month) 

if targetedInstitution == " ": 
        targetedInstitution = "your"
else :
    targetedInstitution = "the " + targetedInstitution



# Ramdomizing Academic Skills and Positive traits

academicSkillsFinal = []

acadSkill1 = random.choice(academicSkills)
academicSkills.remove(acadSkill1)
academicSkillsFinal.append(acadSkill1)
acadSkill2 = random.choice(academicSkills)
academicSkills.remove(acadSkill2)
academicSkillsFinal.append(acadSkill2)
acadSkill3 = random.choice(academicSkills)
academicSkills.remove(acadSkill3)
academicSkillsFinal.append(acadSkill3)

if len(academicSkills) > 0:
    acadSkill4 = random.choice(academicSkills)
    academicSkills.remove(acadSkill4)
    academicSkillsFinal.append(acadSkill4)
    if len(academicSkills)>0:
        acadSkill5 = random.choice(academicSkills)
        academicSkills.remove(acadSkill5)
        academicSkillsFinal.append(acadSkill5)
        if len(academicSkills)>0:
            acadSkill6 = random.choice(academicSkills)
            academicSkills.remove(acadSkill6)
            academicSkillsFinal.append(acadSkill6)


positivePersonalityTraitFinal = []
personalTrait1 = random.choice(positivePersonalityTraits)
positivePersonalityTraits.remove(personalTrait1)
positivePersonalityTraitFinal.append(personalTrait1)
personalTrait2 = random.choice(positivePersonalityTraits)
positivePersonalityTraits.remove(personalTrait2)
positivePersonalityTraitFinal.append(personalTrait2)
personalTrait3 = random.choice(positivePersonalityTraits)
positivePersonalityTraits.remove(personalTrait3)
positivePersonalityTraitFinal.append(personalTrait3)

if len(positivePersonalityTraits)>0:
    personalTrait4 = random.choice(positivePersonalityTraits)
    positivePersonalityTraits.remove(personalTrait4)
    positivePersonalityTraitFinal.append(personalTrait4)
    if len(positivePersonalityTraits)>0:
        personalTrait5 = random.choice(positivePersonalityTraits)
        positivePersonalityTraits.remove(personalTrait5)
        positivePersonalityTraitFinal.append(personalTrait5)
        if len(positivePersonalityTraits)>0:
            personalTrait6 = random.choice(positivePersonalityTraits)
            positivePersonalityTraits.remove(personalTrait6)
            positivePersonalityTraitFinal.append(personalTrait6)



_space = " "

# style = doc.styles.add_style('Indent', WD_STYLE_TYPE.PARAGRAPH)
# paragraph_format = style.paragraph_format
# # paragraph_format.left_indent = Inches(0.25)
# # paragraph_format.first_line_indent = Inches(-0.25)
# paragraph_format.line_spacing_rule = Pt(10)

###############
# 1st Paragraph
###############

doc.add_paragraph("To whom it may concern,\n")

_1st_paragraph = doc.add_paragraph ("\t" + random.choice(Phrase1) + firstName + _space + middleName + _space + lastName + " to " + targetedInstitution + _space + purposeOfTheLetter + ". " + random.choice(Phrase2) + _space + objective.lower() + _space + "in my classroom. " + firstName + " was a " + highSchoolYearAttended + " in my " + classAttended + " class in " + schoolYearAttended + ". ")

if(num_months <= 12) : _1st_paragraph.add_run("Although I have only taught " + firstName + _space + "for " + str(int(num_months)) + " months, I can already see ")
if(num_months > 12 and num_months <= 24) : _1st_paragraph.add_run("I have known " + firstName + _space + "for over an year, and " + subjective.lower() + " made an impression in me due to ")
if(num_months > 24) : _1st_paragraph.add_run("I have known " + firstName + _space + "for more than " + str(int(num_months/12)) + " years, and " + subjective.lower() + " made an impression in me due to ")

_1st_paragraph.add_run(possessive.lower() + _space + random.choice(academicSkillsFinal) + _space + "and also " +  random.choice(positivePersonalityTraitFinal) + " personality. ")

_1st_paragraph.add_run(random.choice(Phrase3) +  objective.lower() + " in this letter, and why " + subjective.lower() +  " deserves to be considered in your institution.")

doc.add_paragraph("")


###############
# 2nd Paragraph
###############

acadSkillA = random.choice(academicSkillsFinal)
academicSkillsFinal.remove(acadSkillA)

acadSkillB = random.choice(academicSkillsFinal)
academicSkillsFinal.remove(acadSkillB)

acadSkillC = random.choice(academicSkillsFinal)
academicSkillsFinal.remove(acadSkillC)

doc.add_paragraph("\t" + "While in class, I have observed some remarkable academic skills. " + firstName + _space + AcademicSkills[acadSkillA] + ". " + subjective  + " also " + AcademicSkills[acadSkillB]
    + random.choice(Phrase5) + subjective.lower() + _space + AcademicSkills[acadSkillC] + ".")

doc.add_paragraph("")


################
#3rd Paragraph
################

personalTraitA = random.choice(positivePersonalityTraitFinal)
positivePersonalityTraitFinal.remove(personalTraitA)

personalTraitB = random.choice(positivePersonalityTraitFinal)
positivePersonalityTraitFinal.remove(personalTraitB)

personalTraitC = random.choice(positivePersonalityTraitFinal)
positivePersonalityTraitFinal.remove(personalTraitC)

doc.add_paragraph("\t" + "Besides all " + possessive.lower() + " Academic work, " + firstName + _space + "is a very " + personalTraitA + " student. " + subjective + _space +  PositivePersonalityTraits[personalTraitB] + ". " + subjective  + " also " + PositivePersonalityTraits[personalTrait2]
+ random.choice(Phrase5) + subjective.lower() + _space + PositivePersonalityTraits[personalTraitC]+ ".")

doc.add_paragraph("")


lastParagraph = doc.add_paragraph("\nPlease, contact me if you have any questions.\n")
lastParagraph.add_run("\nSincerely, \n")

lastParagraph.add_run(teachersName + "\n")
lastParagraph.add_run(str(date.today().month) + "/" + str(date.today().day) + "/" + str(date.today().year))

############################################################

doc.add_paragraph("\n")

response = requests.get(signature)
signature_image = BytesIO(response.content)

doc.add_picture(signature_image, width=Inches(1.0), height=Inches(0.8))
doc.add_paragraph(workingAs)
doc.add_paragraph("( " + professorEmail + " )")

############################################################

file_name_docx = (firstName + lastName + "-" + str(date.today().month) + "-" + str(date.today().day) + "-" + str(date.today().year) + ".docx")

doc.save(file_name_docx)

# Set your Cloudinary credentials
cloudinary.config(
    cloud_name="dvkcyhb6q",
    api_key="612183198313841",
    api_secret="I7N0gitYtpNXxf85G9oT9gPe_O4"
)

# Upload the document to Cloudinary
docx_upload_result = cloudinary.uploader.upload(file_name_docx, resource_type="raw")
docx_url = docx_upload_result["url"]

# Remove the local DOCX file
os.remove(file_name_docx)

# Send the URL in the response
print(docx_url)